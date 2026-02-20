from __future__ import annotations

import io
import os
import re
from dataclasses import dataclass, asdict
from typing import Any, Dict, List, Optional, Tuple

import pdfplumber
from docx import Document
from fastapi import FastAPI, File, UploadFile
from fastapi.responses import HTMLResponse
from fastapi.templating import Jinja2Templates
from starlette.requests import Request
from fastapi.staticfiles import StaticFiles
from fastapi import HTTPException
from fastapi.responses import HTMLResponse, JSONResponse
from fastapi.staticfiles import StaticFiles
from starlette.middleware.cors import CORSMiddleware

# Optional: improved PDF extraction
try:
    import fitz  # PyMuPDF
except Exception:
    fitz = None  # type: ignore


APP_DIR = os.path.dirname(os.path.abspath(__file__))
FRONTEND_DIR = os.path.abspath(os.path.join(APP_DIR, "..", "frontend"))

app = FastAPI(title="ATS Resume Checker", version="3.0.0")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=False,
    allow_methods=["*"],
    allow_headers=["*"],
)

app.mount("/static", StaticFiles(directory=FRONTEND_DIR), name="static")


@app.get("/", response_class=HTMLResponse)
def index() -> str:
    with open(os.path.join(FRONTEND_DIR, "index.html"), "r", encoding="utf-8") as f:
        return f.read()


# ---------------- Extraction ----------------

def extract_text_pdf_pymupdf(pdf_bytes: bytes, max_pages: int = 8) -> Tuple[str, Dict[str, Any]]:
    meta: Dict[str, Any] = {"method": "pymupdf", "pages": 0, "per_page_characters": []}
    if fitz is None:
        return "", {"method": "pymupdf_unavailable", "pages": 0, "per_page_characters": []}
    parts: List[str] = []
    doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    meta["pages"] = doc.page_count
    for i in range(min(max_pages, doc.page_count)):
        page = doc.load_page(i)
        t = page.get_text("text") or ""
        meta["per_page_characters"].append(len(t))
        parts.append(t)
    doc.close()
    return "\n".join(parts).strip(), meta


def extract_text_pdf_pdfplumber(pdf_bytes: bytes, max_pages: int = 8) -> Tuple[str, Dict[str, Any]]:
    meta: Dict[str, Any] = {"method": "pdfplumber", "pages": 0, "per_page_characters": []}
    parts: List[str] = []
    with pdfplumber.open(io.BytesIO(pdf_bytes)) as pdf:
        meta["pages"] = len(pdf.pages)
        for page in pdf.pages[:max_pages]:
            t = page.extract_text() or ""
            meta["per_page_characters"].append(len(t))
            parts.append(t)
    return "\n".join(parts).strip(), meta


def extract_text_from_pdf(pdf_bytes: bytes) -> Tuple[str, Dict[str, Any]]:
    """
    Best-effort extraction:
    1) Try PyMuPDF (often better with some PDFs)
    2) Fallback to pdfplumber
    """
    text1, meta1 = extract_text_pdf_pymupdf(pdf_bytes)
    text2, meta2 = ("", {})
    if len(text1) < 200:  # fallback if too little
        text2, meta2 = extract_text_pdf_pdfplumber(pdf_bytes)
        if len(text2) > len(text1):
            return text2, {"primary": meta2, "fallback": meta1}
    return text1, {"primary": meta1, "fallback": meta2}


def extract_text_from_docx(docx_bytes: bytes) -> Tuple[str, Dict[str, Any]]:
    doc = Document(io.BytesIO(docx_bytes))
    parts: List[str] = []
    for p in doc.paragraphs:
        if p.text and p.text.strip():
            parts.append(p.text.strip())
    # tables
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join([c.text.strip() for c in row.cells if c.text and c.text.strip()])
            if row_text:
                parts.append(row_text)
    text = "\n".join(parts).strip()
    return text, {"method": "python-docx", "paragraphs": len(doc.paragraphs), "tables": len(doc.tables)}


# ---------------- Parsing ----------------

EMAIL_RE = re.compile(r"\b[A-Z0-9._%+-]+@[A-Z0-9.-]+\.[A-Z]{2,}\b", re.IGNORECASE)
PHONE_RE = re.compile(r"(?:(?:\+|00)\d{1,3}[\s\-]?)?(?:\(?\d{2,4}\)?[\s\-]?)?\d{3,4}[\s\-]?\d{3,4}\b")
LINKEDIN_RE = re.compile(r"(https?://)?(www\.)?linkedin\.com/[A-Za-z0-9_/\-%.]+", re.IGNORECASE)
GITHUB_RE = re.compile(r"(https?://)?(www\.)?github\.com/[A-Za-z0-9_\-%.]+", re.IGNORECASE)
PORTFOLIO_RE = re.compile(r"\b(?:https?://)?[A-Za-z0-9.-]+\.[A-Za-z]{2,}(?:/[A-Za-z0-9_/\-%.]+)?\b")

SECTION_ALIASES = {
    "summary": ["summary", "profile", "about", "about me", "professional summary", "career summary", "objective"],
    "skills": ["skills", "technical skills", "core skills", "key skills", "competencies", "tools", "technologies"],
    "experience": ["experience", "work experience", "employment", "professional experience", "work history", "internship"],
    "education": ["education", "academics", "academic background", "qualifications"],
    "projects": ["projects", "personal projects", "selected projects"],
    "certifications": ["certifications", "certificates", "licenses", "training"],
    "languages": ["languages", "language"],
}


def normalize(s: str) -> str:
    return re.sub(r"\s+", " ", s.strip())


def lines_from_text(text: str) -> List[str]:
    return [l.strip() for l in text.splitlines() if l.strip()]


def guess_name(lines: List[str]) -> Optional[str]:
    blacklist = {"curriculum vitae", "resume", "cv"}
    for l in lines[:16]:
        low = l.lower().strip(":-•* ").strip()
        if any(b in low for b in blacklist):
            continue
        if "@" in l:
            continue
        if sum(ch.isdigit() for ch in l) > 0:
            continue
        if len(l) > 50:
            continue
        words = re.findall(r"[A-Za-zÀ-ÖØ-öø-ÿ]+", l)
        if 2 <= len(words) <= 4:
            return normalize(l)
    return None


def find_contacts(text: str) -> Dict[str, Any]:
    emails = sorted(set(EMAIL_RE.findall(text)))
    phones = sorted(set([m.group(0).strip() for m in PHONE_RE.finditer(text)]))
    phones = [p for p in phones if len(re.sub(r"\D", "", p)) >= 9][:3]
    linkedin = sorted(set([m.group(0) for m in LINKEDIN_RE.finditer(text)]))[:2]
    github = sorted(set([m.group(0) for m in GITHUB_RE.finditer(text)]))[:2]

    # Other portfolio links (avoid duplicates)
    other = []
    for m in PORTFOLIO_RE.finditer(text):
        u = m.group(0)
        lu = u.lower()
        if "linkedin.com" in lu or "github.com" in lu:
            continue
        if "@" in u:
            continue
        if len(u) < 10:
            continue
        other.append(u)
    other = sorted(set(other))[:2]

    return {"emails": emails, "phones": phones, "linkedin": linkedin, "github": github, "portfolio": other}


def sectionize(lines: List[str]) -> Dict[str, List[str]]:
    alias_to_key: Dict[str, str] = {}
    for key, aliases in SECTION_ALIASES.items():
        for a in aliases:
            alias_to_key[a.lower()] = key

    def heading_key(line: str) -> Optional[str]:
        l = line.lower().strip(":-•* ").strip()
        if len(l) > 44:
            return None
        return alias_to_key.get(l)

    sections: Dict[str, List[str]] = {k: [] for k in SECTION_ALIASES.keys()}
    sections["_unknown"] = []

    current = "_unknown"
    for line in lines:
        key = heading_key(line)
        if key:
            current = key
            continue
        sections.setdefault(current, []).append(line)
    return sections


def split_skills(skills_lines: List[str]) -> List[str]:
    t = " ".join(skills_lines)
    # Remove category prefixes like "Frontend:"
    t = re.sub(r"\b[A-Za-z][A-Za-z\s/&-]{2,30}:\s*", "", t)
    parts = re.split(r"[•·\u2022\|\n,;/]+", t)
    cleaned = []
    for p in parts:
        p = normalize(p)
        if not p:
            continue
        if len(p) > 45:
            continue
        cleaned.append(p)
    out, seen = [], set()
    for s in cleaned:
        k = s.lower()
        if k in seen:
            continue
        seen.add(k)
        out.append(s)
    return out[:70]


def extract_structured(text: str) -> Dict[str, Any]:
    lines = lines_from_text(text)
    name = guess_name(lines)
    contacts = find_contacts(text)
    sections = sectionize(lines)

    about = ""
    if sections.get("summary"):
        about = normalize(" ".join(sections["summary"])[:1400])

    skills = split_skills(sections.get("skills", []))

    def block(key: str, limit: int) -> str:
        return "\n".join(sections.get(key, [])[:limit]).strip()

    return {
        "name": name,
        "contacts": contacts,
        "about": about,
        "skills": skills,
        "experience": block("experience", 120),
        "projects": block("projects", 120),
        "education": block("education", 90),
        "certifications": block("certifications", 70),
        "languages": block("languages", 30),
        "raw_preview": "\n".join(lines[:90]),
        "sections_detected": {k: len(v) for k, v in sections.items() if k != "_unknown"},
    }


# ---------------- Scoring (dashboard style) ----------------

@dataclass
class CheckItem:
    key: str
    label: str
    score: int          # 0-100
    status: str         # pass | warn | fail
    note: str

@dataclass
class DashboardScore:
    score: int          # 0-100
    issues: int
    ats_parse_rate: int # 0-100
    checks: List[CheckItem]
    ats_friendly: bool
    grade: str
    suggestions: List[str]
    signals: Dict[str, Any]


def clamp(v: float, lo: float, hi: float) -> float:
    return max(lo, min(hi, v))


def compute_dashboard(text: str, meta: Dict[str, Any], structured: Dict[str, Any]) -> DashboardScore:
    # Signals
    words = re.findall(r"\b\w+\b", text)
    word_count = len(words)
    char_count = len(text)

    # meta["primary"] has per_page_characters in our format
    primary = meta.get("primary", {}) if isinstance(meta, dict) else {}
    per_page = primary.get("per_page_characters", []) or []
    pages = int(primary.get("pages", 1) or 1)
    avg_chars = (sum(per_page) / max(1, len(per_page))) if per_page else (char_count / max(1, pages))

    likely_scanned = char_count < 800 or avg_chars < 250

    bullet_count = len(re.findall(r"[•\-\u2022]\s", text))
    date_mentions = len(re.findall(r"\b(20\d{2}|19\d{2})\b", text))
    possible_columns = text.count("    ") > 160

    signals = {
        "pages": pages,
        "word_count": word_count,
        "avg_chars_per_page": round(avg_chars, 1),
        "likely_scanned_pdf": bool(likely_scanned),
        "possible_columns": bool(possible_columns),
        "bullet_count": bullet_count,
        "date_mentions": date_mentions,
        "char_count": char_count,
        "extractor": meta.get("primary", {}).get("method") if isinstance(meta, dict) else None,
        "pymupdf_available": bool(fitz is not None),
    }

    # Checks
    checks: List[CheckItem] = []
    suggestions: List[str] = []
    issues = 0

    # 1) ATS parse rate
    if likely_scanned:
        ats_parse = 30
        issues += 1
        suggestions.append("Export your resume as a text-based PDF (not scanned). If needed, run OCR before uploading.")
        checks.append(CheckItem("ats_parse", "ATS Parse Rate", 30, "fail", "Looks scanned or image-based."))
    else:
        ats_parse = 92
        checks.append(CheckItem("ats_parse", "ATS Parse Rate", 92, "pass", "Text extraction works."))

    # 2) Contact completeness
    emails = structured["contacts"]["emails"]
    phones = structured["contacts"]["phones"]
    has_link = bool(structured["contacts"]["linkedin"] or structured["contacts"]["github"] or structured["contacts"]["portfolio"])
    contact_score = 100
    note_parts = []
    if not emails:
        contact_score -= 45
        issues += 1
        suggestions.append("Add a professional email near the top in plain text.")
        note_parts.append("No email found")
    if not phones:
        contact_score -= 25
        issues += 1
        suggestions.append("Add a phone number in plain text (include country code).")
        note_parts.append("No phone found")
    if not has_link:
        contact_score -= 10
        suggestions.append("Add LinkedIn/GitHub/portfolio link in plain text.")
        note_parts.append("No profile links found")

    contact_score = int(clamp(contact_score, 0, 100))
    status = "pass" if contact_score >= 85 else ("warn" if contact_score >= 60 else "fail")
    if status != "pass":
        issues += 0
    checks.append(CheckItem("contact", "Contact Details", contact_score, status, ", ".join(note_parts) if note_parts else "Looks complete."))

    # 3) Sections check
    sec = structured.get("sections_detected", {})
    has_skills = sec.get("skills", 0) > 0
    has_edu = sec.get("education", 0) > 0
    has_exp_or_proj = (sec.get("experience", 0) > 0) or (sec.get("projects", 0) > 0)
    sec_score = 100
    missing = []
    if not has_skills:
        sec_score -= 35; issues += 1; missing.append("Skills")
        suggestions.append("Add a clear 'Skills' section with keywords (tools, languages, platforms).")
    if not has_edu:
        sec_score -= 20; issues += 1; missing.append("Education")
        suggestions.append("Add an 'Education' section with degree + dates.")
    if not has_exp_or_proj:
        sec_score -= 35; issues += 1; missing.append("Experience/Projects")
        suggestions.append("Add 'Experience' or 'Projects' with impact bullets and tech used.")
    sec_score = int(clamp(sec_score, 0, 100))
    status = "pass" if sec_score >= 85 else ("warn" if sec_score >= 60 else "fail")
    checks.append(CheckItem("sections", "Sections", sec_score, status, ("Missing: " + ", ".join(missing)) if missing else "All key sections detected."))

    # 4) Impact / quantification
    nums = len(re.findall(r"\b\d+%|\b\d+\b", text))
    impact_score = 80
    if nums < 6:
        impact_score -= 25
        issues += 1
        suggestions.append("Add measurable results: %, time saved, users, revenue, speed, KPIs.")
        note = "Few numbers/metrics found."
        status = "warn"
    else:
        note = "Good amount of metrics."
        status = "pass"
        impact_score = 92
    checks.append(CheckItem("impact", "Quantifying Impact", int(clamp(impact_score, 0, 100)), status, note))

    # 5) Repetition
    # Simple repetition heuristic: top repeated words
    low_words = [w.lower() for w in words if len(w) >= 4]
    freq: Dict[str, int] = {}
    for w in low_words:
        freq[w] = freq.get(w, 0) + 1
    top = sorted(freq.items(), key=lambda x: x[1], reverse=True)[:6]
    rep_score = 90
    rep_note = "Looks fine."
    if top and top[0][1] >= 18:
        rep_score = 65
        issues += 1
        rep_note = f"High repetition of '{top[0][0]}' ({top[0][1]}x)."
        suggestions.append("Reduce repeated words; vary action verbs and rewrite duplicated phrases.")
        status = "warn"
    else:
        status = "pass"
    checks.append(CheckItem("repetition", "Repetition", rep_score, status, rep_note))

    # 6) Brevity / length
    brev_score = 92
    brev_note = "Length looks reasonable."
    if word_count < 220:
        brev_score = 62
        issues += 1
        brev_note = "Resume is short (low keyword coverage)."
        suggestions.append("Add more relevant bullets, projects, tools, and responsibilities.")
        status = "warn"
    elif word_count > 1200:
        brev_score = 60
        issues += 1
        brev_note = "Resume is long (harder to scan)."
        suggestions.append("Aim for 1–2 pages and prioritize relevant content.")
        status = "warn"
    else:
        status = "pass"
    checks.append(CheckItem("brevity", "Format & Brevity", brev_score, status, brev_note))

    # 7) Spelling/grammar (light heuristic, offline)
    # We can't do full grammar here without heavy NLP; show as informational.
    grammar_score = 80
    grammar_note = "Basic check only (offline)."
    # Penalize many ALL CAPS words or weird symbols
    caps = len([w for w in re.findall(r"\b[A-Z]{4,}\b", text)])
    weird = len(re.findall(r"[^\w\s•\-\u2022,.;:/()@+%#&]", text))
    if caps > 25 or weird > 80:
        grammar_score = 65
        issues += 1
        grammar_note = "Formatting noise detected (icons/symbols/caps)."
        suggestions.append("Avoid too many icons/symbols; keep headings consistent and readable.")
        status = "warn"
    else:
        status = "pass"
        grammar_score = 86
    checks.append(CheckItem("grammar", "Spelling & Grammar", grammar_score, status, grammar_note))

    # Overall score: weighted
    overall = (
        0.30 * ats_parse +
        0.15 * contact_score +
        0.15 * sec_score +
        0.12 * impact_score +
        0.10 * rep_score +
        0.10 * brev_score +
        0.08 * grammar_score
    )
    overall = int(round(clamp(overall, 0, 100)))

    ats_friendly = (ats_parse >= 70) and (overall >= 60)

    if overall >= 90:
        grade = "Excellent"
    elif overall >= 80:
        grade = "Great"
    elif overall >= 70:
        grade = "Good"
    elif overall >= 60:
        grade = "Fair"
    else:
        grade = "Needs work"

    # Dedupe suggestions
    dedup = []
    seen = set()
    for s in suggestions:
        k = s.lower()
        if k in seen:
            continue
        seen.add(k)
        dedup.append(s)

    return DashboardScore(
        score=overall,
        issues=int(issues),
        ats_parse_rate=int(ats_parse),
        checks=checks,
        ats_friendly=ats_friendly,
        grade=grade,
        suggestions=dedup[:12],
        signals=signals
    )


@app.post("/api/analyze")
async def analyze(file: UploadFile = File(...)) -> JSONResponse:
    name = file.filename or "resume"
    lower = name.lower()

    data = await file.read()
    if not data or len(data) < 200:
        raise HTTPException(status_code=400, detail="Empty or invalid file.")

    text = ""
    meta: Dict[str, Any] = {}

    if lower.endswith(".pdf"):
        text, meta = extract_text_from_pdf(data)
    elif lower.endswith(".docx"):
        text, meta = extract_text_from_docx(data)
    else:
        raise HTTPException(status_code=400, detail="Upload a PDF or DOCX file.")

    structured = extract_structured(text) if text else {
        "name": None,
        "contacts": {"emails": [], "phones": [], "linkedin": [], "github": [], "portfolio": []},
        "about": "",
        "skills": [],
        "experience": "",
        "projects": "",
        "education": "",
        "certifications": "",
        "languages": "",
        "raw_preview": "",
        "sections_detected": {},
    }

    dashboard = compute_dashboard(text, meta, structured) if text else DashboardScore(
        score=20, issues=1, ats_parse_rate=10, checks=[],
        ats_friendly=False, grade="Needs work",
        suggestions=["No text could be extracted. If this is a scanned PDF, run OCR or export as a text-based PDF."],
        signals={"extraction_failed": True, **(meta or {})},
    )

    return JSONResponse(content={
        "filename": name,
        "meta": meta,
        "structured": structured,
        "dashboard": {
            **asdict(dashboard),
            "checks": [asdict(c) for c in dashboard.checks],
        }
    })
